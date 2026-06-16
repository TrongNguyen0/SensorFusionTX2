from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_FILE = "Bao_cao_SensorFusion_Chuong_1_2_3.docx"
TITLE = "Sensor Fusion giữa LiDAR 2D và Camera RGB-D phục vụ bổ sung độ sâu"


TOC_ITEMS = [
    ("CHƯƠNG 1. GIỚI THIỆU", 0),
    ("1.1. Bối cảnh nghiên cứu", 1),
    ("1.2. Lý do chọn đề tài", 1),
    ("1.3. Mục tiêu của đề tài", 1),
    ("1.4. Đối tượng và phạm vi nghiên cứu", 1),
    ("1.5. Cấu trúc báo cáo", 1),
    ("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", 0),
    ("2.1. Tổng quan về Sensor Fusion", 1),
    ("2.2. Mô hình dữ liệu của LiDAR 2D", 1),
    ("2.3. Mô hình camera RGB-D", 1),
    ("2.4. Hệ tọa độ LiDAR và hệ tọa độ camera", 1),
    ("2.5. Phép chiếu điểm 3D lên mặt phẳng ảnh", 1),
    ("2.6. Bài toán PnP trong calibration ngoại tại", 1),
    ("2.7. Vai trò của RANSAC trong loại bỏ nhiễu", 1),
    ("2.8. Sai số tái chiếu", 1),
    ("CHƯƠNG 3. THIẾT KẾ VÀ TRIỂN KHAI HỆ THỐNG", 0),
    ("3.1. Tổng quan kiến trúc hệ thống", 1),
    ("3.2. Thành phần phần cứng", 1),
    ("3.3. Thành phần phần mềm", 1),
    ("3.4. Quy trình thu thập dữ liệu calibration", 1),
    ("3.5. Định dạng dữ liệu calibration", 1),
    ("3.6. Tính toán calibration", 1),
    ("3.7. Kiểm chứng bằng fusion thời gian thực", 1),
    ("3.8. Giới hạn của bản triển khai hiện tại", 1),
    ("CHƯƠNG 4. QUY TRÌNH THU THẬP DỮ LIỆU CALIBRATION", 0),
    ("CHƯƠNG 5. MÔ HÌNH TOÁN HỌC CALIBRATION", 0),
    ("CHƯƠNG 6. GIẢI THUẬT HIỆU CHUẨN PNP/RANSAC", 0),
    ("CHƯƠNG 7. TRIỂN KHAI CHƯƠNG TRÌNH", 0),
    ("CHƯƠNG 8. THIẾT KẾ THỰC NGHIỆM", 0),
    ("CHƯƠNG 9. KẾT QUẢ VÀ ĐÁNH GIÁ", 0),
    ("CHƯƠNG 10. KIỂM CHỨNG BẰNG FUSION THỜI GIAN THỰC", 0),
    ("CHƯƠNG 11. THẢO LUẬN", 0),
    ("CHƯƠNG 12. KẾT LUẬN", 0),
    ("TÀI LIỆU THAM KHẢO", 0),
    ("PHỤ LỤC", 0),
]


REMAINING_HEADINGS = [
    ("CHƯƠNG 4. QUY TRÌNH THU THẬP DỮ LIỆU CALIBRATION", 1),
    ("4.1. Khởi tạo cảm biến", 2),
    ("4.2. Lọc vùng quan sát LiDAR", 2),
    ("4.3. Khử nhiễu dữ liệu LiDAR", 2),
    ("4.4. Chọn điểm target trên dữ liệu LiDAR", 2),
    ("4.5. Chọn điểm target trên ảnh RGB", 2),
    ("4.6. Tạo cặp tương ứng 3D-2D", 2),
    ("4.7. Lưu dữ liệu calibration", 2),
    ("CHƯƠNG 5. MÔ HÌNH TOÁN HỌC CALIBRATION", 1),
    ("5.1. Chuyển đổi tọa độ cực sang tọa độ Cartesian", 2),
    ("5.2. Biểu diễn điểm LiDAR trong không gian 3D", 2),
    ("5.3. Biểu diễn điểm ảnh 2D", 2),
    ("5.4. Biến đổi ngoại tại từ LiDAR sang camera", 2),
    ("5.5. Phép chiếu phối cảnh lên mặt phẳng ảnh", 2),
    ("CHƯƠNG 6. GIẢI THUẬT HIỆU CHUẨN PNP/RANSAC", 1),
    ("6.1. Trích xuất camera intrinsics", 2),
    ("6.2. Nạp dữ liệu pair_*.json", 2),
    ("6.3. Ước lượng ngoại tại bằng solvePnPRansac", 2),
    ("6.4. Chuyển đổi vector quay sang ma trận quay", 2),
    ("6.5. Tính sai số tái chiếu", 2),
    ("6.6. Lưu kết quả calibration", 2),
    ("CHƯƠNG 7. TRIỂN KHAI CHƯƠNG TRÌNH", 1),
    ("7.1. Module collect_calibration.py", 2),
    ("7.2. Module compute_calibration.py", 2),
    ("7.3. Module fusion_calibration.py", 2),
    ("7.4. Cấu trúc dữ liệu đầu vào và đầu ra", 2),
    ("7.5. Xử lý lỗi và giải phóng tài nguyên", 2),
    ("CHƯƠNG 8. THIẾT KẾ THỰC NGHIỆM", 1),
    ("8.1. Kịch bản thu thập dữ liệu", 2),
    ("8.2. Điều kiện môi trường", 2),
    ("8.3. Tiêu chí chất lượng dữ liệu", 2),
    ("8.4. Chỉ số đánh giá", 2),
    ("CHƯƠNG 9. KẾT QUẢ VÀ ĐÁNH GIÁ", 1),
    ("9.1. Kết quả camera intrinsics", 2),
    ("9.2. Kết quả ma trận quay và vector tịnh tiến", 2),
    ("9.3. Thống kê inlier và outlier", 2),
    ("9.4. Đánh giá sai số tái chiếu", 2),
    ("9.5. Minh họa kết quả chiếu điểm LiDAR lên ảnh", 2),
    ("CHƯƠNG 10. KIỂM CHỨNG BẰNG FUSION THỜI GIAN THỰC", 1),
    ("10.1. Quy trình chạy fusion", 2),
    ("10.2. Hiển thị ảnh RGB, depth và overlay", 2),
    ("10.3. Đánh giá định tính kết quả fusion", 2),
    ("10.4. Lưu ảnh kiểm chứng", 2),
    ("CHƯƠNG 11. THẢO LUẬN", 1),
    ("11.1. Nguồn sai số từ LiDAR", 2),
    ("11.2. Nguồn sai số từ camera", 2),
    ("11.3. Sai số do thao tác chọn điểm", 2),
    ("11.4. Giới hạn của LiDAR 2D", 2),
    ("11.5. Đề xuất cải tiến", 2),
    ("CHƯƠNG 12. KẾT LUẬN", 1),
    ("12.1. Kết luận chính", 2),
    ("12.2. Hướng phát triển", 2),
    ("TÀI LIỆU THAM KHẢO", 1),
    ("PHỤ LỤC", 1),
    ("Phụ lục A. Cấu trúc file JSON calibration", 2),
    ("Phụ lục B. Hướng dẫn chạy chương trình", 2),
    ("Phụ lục C. Checklist phần cứng", 2),
]


def set_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(13)
    styles["Normal"].paragraph_format.line_spacing = 1.3
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Times New Roman"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        styles[name].font.bold = True

    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(13)
    return doc


def add_centered(doc, text, size=15, bold=True, after=12):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, size=15 if level == 1 else 14 if level == 2 else 13, bold=True)


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.0)
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, size=13)


def add_formula(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_font(run, size=13, italic=True)


def add_toc_page(doc):
    add_centered(doc, "MỤC LỤC", size=15, after=8)
    for item, level in TOC_ITEMS:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.7 * level)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(item)
        set_font(run, size=11, bold=level == 0)
    doc.add_page_break()


def add_remaining_headings(doc):
    for text, level in REMAINING_HEADINGS:
        add_heading(doc, text, level)


def build_document():
    doc = setup_document()

    add_centered(doc, "BÁO CÁO ĐỀ TÀI", size=16)
    add_centered(doc, TITLE.upper(), size=15)
    doc.add_page_break()

    add_toc_page(doc)

    add_paragraph(
        doc,
        "Tài liệu này trình bày bản thảo ban đầu cho Chương 1, Chương 2 và Chương 3 của báo cáo. Nội dung được xây dựng dựa trên cấu trúc chương trình hiện có của dự án, bao gồm các tệp mã nguồn chính: collect_calibration.py, compute_calibration.py và fusion_calibration.py. Các thông số định lượng chưa được xác nhận từ kết quả thực nghiệm hoặc tài liệu kỹ thuật chính thức được trình bày ở mức khái quát nhằm tránh đưa ra kết luận không có cơ sở."
    )

    add_heading(doc, "CHƯƠNG 1. GIỚI THIỆU", 1)

    add_heading(doc, "1.1. Bối cảnh nghiên cứu", 2)
    add_paragraph(
        doc,
        "Trong các hệ thống nhận thức môi trường, thông tin hình học và thông tin thị giác thường cần được kết hợp để mô tả không gian xung quanh một cách đầy đủ hơn. Camera cung cấp ảnh có mật độ điểm ảnh cao, chứa nhiều thông tin về màu sắc, biên dạng và ngữ cảnh của vật thể. Tuy nhiên, việc suy luận khoảng cách từ ảnh camera đơn lẻ hoặc từ ảnh RGB không phải lúc nào cũng ổn định, đặc biệt trong các điều kiện ánh sáng phức tạp hoặc khi bề mặt vật thể thiếu đặc trưng thị giác."
    )
    add_paragraph(
        doc,
        "LiDAR là cảm biến có khả năng đo khoảng cách trực tiếp dựa trên nguyên lý phát và thu tín hiệu laser. Đối với LiDAR 2D, dữ liệu thu được thường có dạng tập hợp các cặp góc quét và khoảng cách. Dữ liệu này có độ tin cậy nhất định về mặt hình học trong mặt phẳng quét, nhưng lại thưa hơn ảnh camera và không chứa thông tin màu sắc. Do đó, việc kết hợp LiDAR 2D với camera RGB-D là một hướng tiếp cận phù hợp để tận dụng ưu điểm của từng loại cảm biến."
    )
    add_paragraph(
        doc,
        "Trong đề tài này, hệ thống Sensor Fusion được xây dựng với mục tiêu hiệu chuẩn mối quan hệ hình học giữa LiDAR 2D và camera RGB-D. Sau khi hiệu chuẩn, các điểm đo từ LiDAR có thể được chiếu lên mặt phẳng ảnh camera. Kết quả này là cơ sở để kiểm chứng sự tương ứng giữa dữ liệu khoảng cách của LiDAR và hình ảnh thu được từ camera, đồng thời phục vụ bài toán bổ sung hoặc kiểm chứng thông tin độ sâu."
    )

    add_heading(doc, "1.2. Lý do chọn đề tài", 2)
    add_paragraph(
        doc,
        "Bài toán bổ sung độ sâu từ nhiều cảm biến có ý nghĩa thực tiễn trong các hệ thống robot di động, xe tự hành mô hình, giám sát không gian và các ứng dụng đo lường trong môi trường trong nhà. Camera RGB-D có thể cung cấp ảnh màu và ảnh độ sâu, trong khi LiDAR 2D cung cấp các điểm đo khoảng cách theo góc quét. Khi hai nguồn dữ liệu này được đặt trong cùng một hệ quy chiếu, hệ thống có thể khai thác đồng thời thông tin ảnh và thông tin khoảng cách."
    )
    add_paragraph(
        doc,
        "Một yêu cầu quan trọng của quá trình kết hợp cảm biến là calibration, tức xác định quan hệ hình học giữa các hệ tọa độ cảm biến. Nếu calibration không chính xác, điểm LiDAR sau khi chiếu lên ảnh sẽ bị lệch khỏi vị trí vật thể thực tế, dẫn đến sai số trong quá trình phân tích hoặc bổ sung độ sâu. Vì vậy, phần calibration giữ vai trò nền tảng trong toàn bộ pipeline Sensor Fusion."
    )

    add_heading(doc, "1.3. Mục tiêu của đề tài", 2)
    add_paragraph(
        doc,
        "Mục tiêu chính của đề tài là xây dựng quy trình hiệu chuẩn giữa LiDAR 2D và camera RGB-D, từ đó phục vụ việc chiếu điểm LiDAR lên ảnh camera và kiểm chứng khả năng kết hợp dữ liệu giữa hai cảm biến. Quy trình này bao gồm thu thập dữ liệu tương ứng giữa điểm LiDAR và điểm ảnh, ước lượng ma trận biến đổi ngoại tại, đánh giá sai số tái chiếu và ứng dụng kết quả hiệu chuẩn vào chương trình fusion thời gian thực."
    )
    add_paragraph(
        doc,
        "Cụ thể, đề tài hướng đến các mục tiêu sau: xây dựng chương trình thu dữ liệu calibration từ RPLidar và Intel RealSense; biểu diễn điểm LiDAR trong hệ tọa độ Cartesian; thiết lập các cặp tương ứng giữa điểm 3D của LiDAR và điểm 2D trên ảnh; sử dụng phương pháp PnP kết hợp RANSAC để ước lượng ma trận quay và vector tịnh tiến; sử dụng kết quả calibration để chiếu điểm LiDAR lên ảnh RGB và kiểm chứng trực quan."
    )

    add_heading(doc, "1.4. Đối tượng và phạm vi nghiên cứu", 2)
    add_paragraph(
        doc,
        "Đối tượng nghiên cứu của đề tài là hệ thống fusion giữa LiDAR 2D RPLidar A1M8 và camera RGB-D Intel RealSense D435. Trong phạm vi hiện tại, LiDAR được xem như cảm biến đo khoảng cách trong mặt phẳng quét, còn camera cung cấp ảnh màu và ảnh độ sâu được căn chỉnh theo khung ảnh màu. Phần calibration tập trung vào việc xác định quan hệ từ hệ tọa độ LiDAR sang hệ tọa độ camera."
    )
    add_paragraph(
        doc,
        "Phạm vi thực nghiệm của đề tài được giới hạn ở môi trường trong phòng, với target calibration dạng thanh phẳng được quan sát đồng thời bởi LiDAR và camera. Việc chọn điểm target trong chương trình hiện tại có sự tham gia của người dùng, bao gồm chọn hai đầu thanh trên biểu đồ LiDAR và chọn hai đầu thanh tương ứng trên ảnh RGB. Các bước tự động nhận dạng target chưa thuộc phạm vi của bản triển khai hiện tại."
    )

    add_heading(doc, "1.5. Cấu trúc báo cáo", 2)
    add_paragraph(
        doc,
        "Báo cáo được tổ chức theo trình tự từ cơ sở lý thuyết đến triển khai hệ thống và đánh giá thực nghiệm. Chương 1 trình bày bối cảnh, lý do chọn đề tài, mục tiêu và phạm vi nghiên cứu. Chương 2 trình bày cơ sở lý thuyết liên quan đến mô hình LiDAR 2D, mô hình camera, phép biến đổi tọa độ, phép chiếu phối cảnh và phương pháp PnP/RANSAC. Chương 3 mô tả thiết kế hệ thống, bao gồm phần cứng, phần mềm, cấu trúc chương trình và quy trình dữ liệu calibration."
    )

    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", 1)

    add_heading(doc, "2.1. Tổng quan về Sensor Fusion", 2)
    add_paragraph(
        doc,
        "Sensor Fusion là quá trình kết hợp dữ liệu từ nhiều cảm biến nhằm tạo ra biểu diễn thông tin đầy đủ, ổn định hoặc có độ tin cậy cao hơn so với việc sử dụng từng cảm biến riêng lẻ. Trong bài toán nhận thức môi trường, mỗi cảm biến thường có ưu điểm và hạn chế riêng. Camera có khả năng ghi nhận chi tiết thị giác của cảnh, trong khi LiDAR cung cấp thông tin khoảng cách trực tiếp theo các hướng quét."
    )
    add_paragraph(
        doc,
        "Đối với hệ thống kết hợp LiDAR 2D và camera RGB-D, fusion không chỉ là ghép dữ liệu ở mức hiển thị mà còn yêu cầu sự thống nhất về mặt hình học. Nói cách khác, một điểm đo trong hệ tọa độ LiDAR cần được chuyển đổi chính xác sang hệ tọa độ camera trước khi có thể chiếu lên ảnh. Quá trình xác định mối quan hệ này được gọi là calibration ngoại tại giữa hai cảm biến."
    )

    add_heading(doc, "2.2. Mô hình dữ liệu của LiDAR 2D", 2)
    add_paragraph(
        doc,
        "LiDAR 2D thu nhận dữ liệu dưới dạng các phép đo khoảng cách theo từng góc quét. Mỗi điểm đo có thể được biểu diễn bằng cặp tọa độ cực gồm góc quét và khoảng cách. Trong chương trình hiện tại, dữ liệu LiDAR được đọc từ RPLidar, sau đó chuẩn hóa góc về miền quan sát phía trước. Các điểm có khoảng cách bằng không hoặc nằm ngoài vùng góc quan tâm sẽ bị loại bỏ."
    )
    add_paragraph(
        doc,
        "Để sử dụng dữ liệu LiDAR trong bài toán calibration, tọa độ cực được chuyển sang tọa độ Cartesian trên mặt phẳng quét. Nếu khoảng cách đo được là d và góc quét là theta, tọa độ của điểm LiDAR được biểu diễn bởi hai thành phần x và z. Trục x biểu diễn phương ngang, còn trục z biểu diễn phương phía trước của cảm biến trong mặt phẳng quét."
    )
    add_formula(doc, "x = d sin(theta),    z = d cos(theta)")
    add_paragraph(
        doc,
        "Do LiDAR sử dụng trong hệ thống là LiDAR 2D, điểm đo được giả định nằm trên một mặt phẳng quét. Vì vậy, khi đưa vào bài toán PnP, điểm LiDAR được biểu diễn dưới dạng điểm 3D với thành phần y bằng không."
    )
    add_formula(doc, "P_L = [x, 0, z]^T")

    add_heading(doc, "2.3. Mô hình camera RGB-D", 2)
    add_paragraph(
        doc,
        "Camera RGB-D cung cấp đồng thời ảnh màu và ảnh độ sâu. Trong hệ thống này, Intel RealSense D435 được cấu hình để thu ảnh màu và ảnh độ sâu ở cùng kích thước khung hình trong chương trình. Dữ liệu depth được căn chỉnh theo khung ảnh màu để thuận tiện cho việc hiển thị và kiểm chứng kết quả fusion."
    )
    add_paragraph(
        doc,
        "Mô hình chiếu camera được sử dụng là mô hình pinhole. Theo mô hình này, một điểm 3D trong hệ tọa độ camera được chiếu lên mặt phẳng ảnh thông qua ma trận nội tại K. Ma trận nội tại chứa các tham số tiêu cự theo hai phương ảnh và tọa độ điểm chính của camera. Trong chương trình compute_calibration.py, các tham số nội tại được lấy trực tiếp từ RealSense SDK tại thời điểm chạy."
    )
    add_formula(doc, "K = [[fx, 0, cx], [0, fy, cy], [0, 0, 1]]")

    add_heading(doc, "2.4. Hệ tọa độ LiDAR và hệ tọa độ camera", 2)
    add_paragraph(
        doc,
        "Do LiDAR và camera được gắn tại hai vị trí vật lý khác nhau, mỗi cảm biến có một hệ tọa độ riêng. Điểm đo trong hệ tọa độ LiDAR không thể trực tiếp sử dụng như điểm trong hệ tọa độ camera. Cần xác định phép biến đổi hình học giữa hai hệ tọa độ này."
    )
    add_paragraph(
        doc,
        "Phép biến đổi từ hệ LiDAR sang hệ camera được biểu diễn bằng ma trận quay R và vector tịnh tiến T. Với một điểm LiDAR P_L, điểm tương ứng trong hệ tọa độ camera P_C được xác định theo công thức:"
    )
    add_formula(doc, "P_C = R P_L + T")
    add_paragraph(
        doc,
        "Trong đó, R mô tả sự khác biệt về hướng giữa hai hệ tọa độ, còn T mô tả độ lệch vị trí giữa gốc tọa độ LiDAR và gốc tọa độ camera. Việc ước lượng chính xác R và T là mục tiêu trung tâm của calibration ngoại tại."
    )

    add_heading(doc, "2.5. Phép chiếu điểm 3D lên mặt phẳng ảnh", 2)
    add_paragraph(
        doc,
        "Sau khi một điểm LiDAR được chuyển sang hệ tọa độ camera, điểm này có thể được chiếu lên ảnh bằng mô hình camera pinhole. Giả sử điểm trong hệ camera có tọa độ P_C = [X_C, Y_C, Z_C]^T, với Z_C là độ sâu theo trục nhìn của camera. Tọa độ pixel tương ứng được tính bằng các tham số nội tại của camera."
    )
    add_formula(doc, "u = fx X_C / Z_C + cx")
    add_formula(doc, "v = fy Y_C / Z_C + cy")
    add_paragraph(
        doc,
        "Điểm chỉ được xem là hợp lệ khi Z_C lớn hơn không và tọa độ pixel nằm bên trong kích thước ảnh. Trong chương trình fusion_calibration.py, các điểm LiDAR hợp lệ được vẽ trực tiếp lên ảnh RGB để tạo ảnh overlay phục vụ kiểm chứng trực quan."
    )

    add_heading(doc, "2.6. Bài toán PnP trong calibration ngoại tại", 2)
    add_paragraph(
        doc,
        "Perspective-n-Point, thường gọi là PnP, là bài toán ước lượng tư thế của camera hoặc quan hệ hình học giữa hệ tọa độ 3D và mặt phẳng ảnh 2D dựa trên các cặp điểm tương ứng. Trong đề tài này, tập điểm 3D được tạo từ dữ liệu LiDAR sau khi chuyển sang dạng [x, 0, z]^T, còn tập điểm 2D là các pixel tương ứng trên ảnh RGB."
    )
    add_paragraph(
        doc,
        "Chương trình compute_calibration.py sử dụng hàm solvePnPRansac của OpenCV để ước lượng vector quay và vector tịnh tiến từ các cặp điểm 3D-2D. Vector quay sau đó được chuyển thành ma trận quay bằng công thức Rodrigues. Kết quả cuối cùng gồm ma trận nội tại K, ma trận quay R và vector tịnh tiến T, được lưu vào tệp calibration_result_pnp.npz để sử dụng trong giai đoạn fusion."
    )

    add_heading(doc, "2.7. Vai trò của RANSAC trong loại bỏ nhiễu", 2)
    add_paragraph(
        doc,
        "Trong quá trình thu thập dữ liệu calibration, sai số có thể xuất hiện do nhiễu đo khoảng cách của LiDAR, sai lệch khi chọn điểm trên ảnh, sai lệch do target không nằm hoàn toàn trong mặt phẳng quét hoặc do thao tác nội suy điểm tương ứng. Những yếu tố này có thể tạo ra các cặp điểm không phù hợp, còn gọi là outlier."
    )
    add_paragraph(
        doc,
        "RANSAC là phương pháp ước lượng tham số có khả năng giảm ảnh hưởng của outlier. Thay vì sử dụng toàn bộ dữ liệu một cách trực tiếp, RANSAC lặp lại quá trình chọn mẫu con, ước lượng mô hình và đánh giá số lượng điểm phù hợp với mô hình. Trong bối cảnh calibration, các điểm có sai số tái chiếu nhỏ hơn ngưỡng đặt trước được xem là inlier và được dùng để đánh giá kết quả."
    )

    add_heading(doc, "2.8. Sai số tái chiếu", 2)
    add_paragraph(
        doc,
        "Sai số tái chiếu là thước đo quan trọng để đánh giá chất lượng calibration. Sau khi ước lượng R và T, mỗi điểm LiDAR 3D được chiếu lên mặt phẳng ảnh. Sai số tái chiếu được tính bằng khoảng cách Euclidean giữa pixel dự đoán và pixel tương ứng được gán trong dữ liệu calibration."
    )
    add_formula(doc, "e_i = || p_i - p'_i ||")
    add_paragraph(
        doc,
        "Trong đó, p_i là điểm ảnh thực nghiệm được chọn hoặc nội suy từ dữ liệu calibration, còn p'_i là điểm ảnh thu được sau khi chiếu điểm LiDAR bằng mô hình đã hiệu chuẩn. Các đại lượng như sai số trung bình, trung vị và sai số lớn nhất có thể được sử dụng để đánh giá mức độ phù hợp của mô hình."
    )

    add_heading(doc, "CHƯƠNG 3. THIẾT KẾ VÀ TRIỂN KHAI HỆ THỐNG", 1)

    add_heading(doc, "3.1. Tổng quan kiến trúc hệ thống", 2)
    add_paragraph(
        doc,
        "Hệ thống được thiết kế theo pipeline gồm ba giai đoạn chính: thu thập dữ liệu calibration, tính toán tham số calibration và kiểm chứng kết quả bằng fusion thời gian thực. Giai đoạn thứ nhất tạo ra các cặp dữ liệu tương ứng giữa điểm LiDAR và điểm ảnh. Giai đoạn thứ hai sử dụng các cặp tương ứng này để ước lượng quan hệ hình học giữa LiDAR và camera. Giai đoạn thứ ba sử dụng kết quả calibration để chiếu điểm LiDAR lên ảnh camera."
    )
    add_paragraph(
        doc,
        "Ba giai đoạn trên tương ứng với ba tệp mã nguồn chính trong thư mục src. Tệp collect_calibration.py đảm nhiệm việc thu dữ liệu; tệp compute_calibration.py đảm nhiệm việc ước lượng ma trận calibration; tệp fusion_calibration.py đảm nhiệm việc kiểm chứng trực quan bằng cách hiển thị ảnh RGB có phủ điểm LiDAR cùng với ảnh depth colormap."
    )

    add_heading(doc, "3.2. Thành phần phần cứng", 2)
    add_heading(doc, "3.2.1. LiDAR 2D RPLidar A1M8", 3)
    add_paragraph(
        doc,
        "LiDAR được sử dụng trong hệ thống là RPLidar A1M8. Trong chương trình, cảm biến được kết nối thông qua cổng serial, với cổng mặc định được đặt là COM3. Dữ liệu LiDAR được đọc dưới dạng các bộ giá trị gồm chất lượng điểm đo, góc quét và khoảng cách. Các điểm được xử lý chủ yếu trong miền góc phía trước của cảm biến để phù hợp với vùng quan sát của camera."
    )
    add_heading(doc, "3.2.2. Camera RGB-D Intel RealSense D435", 3)
    add_paragraph(
        doc,
        "Camera được sử dụng là Intel RealSense D435. Trong chương trình, camera được cấu hình để thu đồng thời luồng ảnh độ sâu và luồng ảnh màu. Các frame depth được căn chỉnh theo frame màu thông qua đối tượng align của RealSense SDK. Cách tổ chức này giúp việc hiển thị ảnh depth và ảnh màu trên cùng hệ tọa độ ảnh trở nên thuận tiện hơn."
    )
    add_heading(doc, "3.2.3. Target calibration", 3)
    add_paragraph(
        doc,
        "Target calibration trong hệ thống là một thanh phẳng có thể được quan sát đồng thời bởi LiDAR và camera. Khi thu dữ liệu, người dùng chọn hai đầu thanh trên biểu đồ LiDAR và hai đầu tương ứng trên ảnh RGB. Các điểm LiDAR nằm trên đoạn thanh được nội suy sang đoạn thẳng trên ảnh, từ đó tạo tập điểm tương ứng phục vụ PnP."
    )
    add_heading(doc, "3.2.4. Bố trí cảm biến", 3)
    add_paragraph(
        doc,
        "LiDAR và camera được lắp đặt sao cho vùng quan sát phía trước của LiDAR nằm trong vùng nhìn của camera. Bố trí cơ khí cần giữ ổn định trong suốt quá trình thu thập dữ liệu và kiểm chứng. Nếu vị trí tương đối giữa hai cảm biến thay đổi sau khi calibration, các tham số R và T đã ước lượng sẽ không còn phản ánh đúng hệ thống thực tế."
    )

    add_heading(doc, "3.3. Thành phần phần mềm", 2)
    add_paragraph(
        doc,
        "Hệ thống được triển khai bằng ngôn ngữ Python. Các thư viện chính bao gồm NumPy cho xử lý mảng và tính toán ma trận, OpenCV cho xử lý ảnh và giải bài toán PnP, pyrealsense2 cho giao tiếp với camera Intel RealSense, rplidar cho giao tiếp với LiDAR, Matplotlib cho hiển thị dữ liệu polar và keyboard cho nhận sự kiện phím trong quá trình thu dữ liệu."
    )
    add_paragraph(
        doc,
        "Cấu trúc thư mục của dự án gồm thư mục src chứa mã nguồn, thư mục data/captured_data chứa dữ liệu thu thập calibration, thư mục data/fusion_output chứa ảnh kết quả fusion, và tệp calibration_result_pnp.npz chứa kết quả calibration sau khi chạy chương trình tính toán."
    )

    add_heading(doc, "3.4. Quy trình thu thập dữ liệu calibration", 2)
    add_paragraph(
        doc,
        "Quy trình thu thập dữ liệu bắt đầu bằng việc khởi tạo LiDAR và camera. LiDAR được đọc liên tục trong một luồng nền để cập nhật dữ liệu scan mới nhất. Camera RealSense được khởi tạo với luồng ảnh màu và ảnh độ sâu, sau đó depth frame được căn chỉnh theo color frame."
    )
    add_paragraph(
        doc,
        "Dữ liệu LiDAR được lọc theo góc quan sát. Các góc lớn hơn 180 độ được chuyển về miền âm tương ứng để tạo miền góc từ âm đến dương. Chương trình chỉ giữ lại các điểm trong khoảng từ -90 độ đến 90 độ và có khoảng cách hợp lệ. Dữ liệu này được dùng để hiển thị trên biểu đồ polar và phục vụ bước chọn target."
    )
    add_paragraph(
        doc,
        "Để giảm ảnh hưởng của nhiễu tức thời, chương trình tích lũy nhiều scan gần nhất, làm tròn góc về giá trị nguyên và lấy trung vị khoảng cách tại mỗi góc. Cách làm này giúp dữ liệu target trên biểu đồ LiDAR ổn định hơn trước khi người dùng chọn hai đầu thanh."
    )
    add_paragraph(
        doc,
        "Sau khi dữ liệu LiDAR được khử nhiễu, người dùng chọn hai đầu thanh trên biểu đồ polar. Chương trình xác định dải điểm LiDAR nằm giữa hai góc được chọn và tiếp tục lọc các điểm có biến thiên khoảng cách bất thường. Tiếp theo, người dùng chọn hai đầu thanh tương ứng trên ảnh RGB. Từ hai cặp đầu mút này, các điểm LiDAR trên thanh được nội suy tuyến tính thành các điểm ảnh tương ứng."
    )

    add_heading(doc, "3.5. Định dạng dữ liệu calibration", 2)
    add_paragraph(
        doc,
        "Mỗi lần thu dữ liệu calibration, chương trình lưu ảnh màu, ảnh độ sâu, ảnh depth colormap, ảnh kết quả hiển thị điểm nội suy và một tệp JSON mô tả dữ liệu tương ứng. Tệp JSON chứa thời điểm thu dữ liệu, hai điểm đầu mút trên ảnh, khoảng góc LiDAR của target, số lượng điểm LiDAR, khoảng cách trung bình, danh sách điểm Cartesian và danh sách mapped_points."
    )
    add_paragraph(
        doc,
        "Trường mapped_points có vai trò quan trọng trong bước tính calibration. Mỗi phần tử trong trường này chứa góc LiDAR, khoảng cách LiDAR và tọa độ pixel tương ứng trên ảnh. Chương trình compute_calibration.py đọc các tệp pair_*.json trong thư mục data/captured_data/pair để tạo tập điểm 3D-2D phục vụ PnP."
    )

    add_heading(doc, "3.6. Tính toán calibration", 2)
    add_paragraph(
        doc,
        "Trong giai đoạn tính toán calibration, chương trình trước hết lấy ma trận nội tại của camera từ RealSense SDK. Sau đó, các tệp JSON đã thu được được nạp vào bộ nhớ. Mỗi điểm mapped_points được chuyển từ dạng polar của LiDAR sang dạng Cartesian, rồi biểu diễn thành điểm 3D với thành phần y bằng không. Pixel tương ứng được đưa vào tập điểm ảnh 2D."
    )
    add_paragraph(
        doc,
        "Tập điểm 3D-2D được truyền vào hàm solvePnPRansac của OpenCV. Hàm này trả về vector quay, vector tịnh tiến và danh sách inlier nếu quá trình ước lượng thành công. Vector quay được chuyển thành ma trận quay R bằng phép biến đổi Rodrigues. Kết quả gồm K, R và T được lưu lại trong tệp calibration_result_pnp.npz."
    )

    add_heading(doc, "3.7. Kiểm chứng bằng fusion thời gian thực", 2)
    add_paragraph(
        doc,
        "Sau khi có kết quả calibration, chương trình fusion_calibration.py được sử dụng để kiểm chứng. Chương trình tải K, R và T từ tệp calibration_result_pnp.npz, đọc frame từ camera và đọc scan từ LiDAR. Mỗi điểm LiDAR hợp lệ được chuyển sang tọa độ Cartesian, biến đổi sang hệ tọa độ camera và chiếu lên mặt phẳng ảnh."
    )
    add_paragraph(
        doc,
        "Các điểm chiếu hợp lệ được vẽ lên ảnh RGB dưới dạng các điểm màu. Ảnh overlay được ghép cạnh ảnh depth colormap để người dùng quan sát đồng thời kết quả chiếu LiDAR và thông tin độ sâu từ camera. Nếu điểm LiDAR nằm gần đúng vị trí vật thể thực tế trên ảnh, kết quả này cho thấy calibration có sự phù hợp về mặt hình học."
    )

    add_heading(doc, "3.8. Giới hạn của bản triển khai hiện tại", 2)
    add_paragraph(
        doc,
        "Bản triển khai hiện tại vẫn phụ thuộc vào thao tác chọn điểm thủ công của người dùng. Do đó, độ chính xác của dữ liệu calibration có thể bị ảnh hưởng bởi sai lệch khi chọn hai đầu target trên biểu đồ LiDAR và ảnh RGB. Ngoài ra, do LiDAR là cảm biến 2D, mô hình điểm 3D sử dụng giả định thành phần y bằng không. Giả định này phù hợp khi target nằm trong mặt phẳng quét của LiDAR, nhưng có thể gây sai số nếu điều kiện bố trí không đảm bảo."
    )
    add_paragraph(
        doc,
        "Các thông số đánh giá định lượng như sai số trung bình, sai số trung vị, sai số lớn nhất và tỷ lệ inlier cần được trình bày dựa trên kết quả chạy thực nghiệm cụ thể. Trong bản thảo này, các kết luận định lượng chưa được đưa ra vì chưa có yêu cầu tổng hợp số liệu từ tập dữ liệu hiện tại."
    )

    add_remaining_headings(doc)

    doc.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
