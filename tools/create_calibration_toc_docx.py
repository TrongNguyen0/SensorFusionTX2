from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_FILE = "Bao_cao_Calibration_Muc_luc.docx"


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph(paragraph, align=None, before=0, after=6, line_spacing=1.15):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, after=0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph()
    style_name = f"Heading {level}"
    paragraph.style = style_name
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 14, bold=True)
    set_paragraph(paragraph, before=8 if level == 1 else 4, after=4)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    set_paragraph(title, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    run = title.add_run("ĐỀ CƯƠNG BÁO CÁO KHOA HỌC")
    set_run_font(run, size=16, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    run = subtitle.add_run(
        "Hiệu chuẩn và kiểm chứng hệ thống Sensor Fusion giữa RPLidar A1M8 và Intel RealSense D435"
    )
    set_run_font(run, size=14, bold=True)

    meta = [
        ("Tên dự án", "SensorFusion - LiDAR Camera Calibration"),
        ("Phạm vi tài liệu", "Mục lục chi tiết cho báo cáo calibration"),
        ("Dữ liệu/code chính", "src/collect_calibration.py, src/compute_calibration.py, src/fusion_calibration.py"),
        ("Kết quả calibration", "calibration_result_pnp.npz"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Thông tin", bold=True)
    set_cell_text(table.rows[0].cells[1], "Nội dung", bold=True)
    shade_cell(table.rows[0].cells[0], "D9EAF7")
    shade_cell(table.rows[0].cells[1], "D9EAF7")
    for key, value in meta:
        row = table.add_row().cells
        set_cell_text(row[0], key, bold=True)
        set_cell_text(row[1], value)

    doc.add_paragraph()

    intro = doc.add_paragraph()
    set_paragraph(intro, after=10)
    run = intro.add_run(
        "Tài liệu này mới ở mức mục lục/khung báo cáo. Các mục được thiết kế để có thể mở rộng thành báo cáo khoa học hoàn chỉnh, bao quát từ phần cứng, phần mềm, quy trình thu thập dữ liệu, mô hình toán học, giải thuật calibration đến đánh giá thực nghiệm."
    )
    set_run_font(run, size=12)

    add_heading(doc, "MỤC LỤC ĐỀ XUẤT", 1)

    toc = [
        ("1", "Giới thiệu", "Bối cảnh, nhu cầu sensor fusion, mục tiêu hiệu chuẩn LiDAR-camera, phạm vi nghiên cứu."),
        ("1.1", "Bối cảnh và động lực nghiên cứu", "Vai trò của LiDAR và camera RGB-D trong nhận thức môi trường."),
        ("1.2", "Bài toán calibration LiDAR-camera", "Định nghĩa bài toán ánh xạ điểm LiDAR vào mặt phẳng ảnh camera."),
        ("1.3", "Mục tiêu và đóng góp của hệ thống", "Xây dựng pipeline thu dữ liệu, tính ngoại tại và kiểm chứng realtime."),
        ("1.4", "Phạm vi và giới hạn", "RPLidar 2D, RealSense D435, môi trường trong phòng, target dạng thanh phẳng."),
        ("2", "Tổng quan cơ sở lý thuyết", "Các khái niệm nền tảng phục vụ calibration và projection."),
        ("2.1", "Mô hình cảm biến LiDAR 2D", "Tọa độ cực, khoảng cách, góc quét, miền quan sát phía trước."),
        ("2.2", "Mô hình camera pinhole và RGB-D", "Ma trận nội tại K, hệ tọa độ camera, ảnh màu và ảnh độ sâu."),
        ("2.3", "Biến đổi hệ tọa độ", "Chuyển từ hệ LiDAR sang hệ camera bằng R và T."),
        ("2.4", "Phép chiếu phối cảnh", "Từ điểm 3D trong camera sang pixel ảnh 2D."),
        ("2.5", "PnP và RANSAC", "Ước lượng ngoại tại từ cặp điểm 3D-2D và loại nhiễu/outlier."),
        ("2.6", "Sai số tái chiếu", "Định nghĩa reprojection error và ý nghĩa trong đánh giá calibration."),
        ("3", "Thiết kế hệ thống phần cứng", "Mô tả đầy đủ thiết bị, cách lắp đặt và điều kiện vận hành."),
        ("3.1", "RPLidar A1M8", "Thông số chính, kiểu dữ liệu, cổng COM, đặc tính quét 2D."),
        ("3.2", "Intel RealSense D435", "Ảnh RGB, depth frame, alignment depth-color, độ phân giải sử dụng."),
        ("3.3", "Máy tính và kết nối", "USB, driver, hệ điều hành, yêu cầu hiệu năng tối thiểu."),
        ("3.4", "Target calibration", "Thanh phẳng phản xạ tốt, cách đặt target, lý do chọn target."),
        ("3.5", "Bố trí thực nghiệm", "Vị trí tương đối LiDAR-camera, khoảng cách đo, vùng quan sát."),
        ("4", "Môi trường phần mềm", "Ngôn ngữ, thư viện, cấu trúc project và dữ liệu đầu vào/đầu ra."),
        ("4.1", "Ngôn ngữ và hệ điều hành", "Python trên Windows, cấu hình COM cho LiDAR."),
        ("4.2", "Thư viện sử dụng", "numpy, OpenCV, pyrealsense2, rplidar, matplotlib, keyboard."),
        ("4.3", "Cấu trúc thư mục project", "src, data/captured_data, data/fusion_output, file calibration_result_pnp.npz."),
        ("4.4", "Các script chính", "collect_calibration.py, compute_calibration.py, fusion_calibration.py."),
        ("4.5", "Định dạng dữ liệu lưu trữ", "Ảnh PNG, depth raw, depth colormap, JSON pair_*.json."),
        ("5", "Quy trình thu thập dữ liệu calibration", "Pipeline tạo cặp tương ứng LiDAR-camera."),
        ("5.1", "Khởi tạo cảm biến", "Mở RealSense stream, align depth-color, kết nối RPLidar COM3."),
        ("5.2", "Lọc vùng quan sát LiDAR", "Chuẩn hóa góc về [-180, 180], giữ miền [-90, 90]."),
        ("5.3", "Khử nhiễu dữ liệu LiDAR", "Gom nhiều scan, làm tròn góc, lấy median khoảng cách."),
        ("5.4", "Chọn điểm target trên polar plot", "Người dùng chọn hai đầu thanh trên dữ liệu LiDAR."),
        ("5.5", "Chọn điểm target trên ảnh RGB", "Người dùng chọn hai đầu thanh tương ứng trên ảnh camera."),
        ("5.6", "Nội suy cặp điểm 3D-2D", "Ánh xạ các điểm trên thanh LiDAR sang đường nối hai pixel ảnh."),
        ("5.7", "Lưu dữ liệu calibration", "Lưu mapped_points, cartesian_points, ảnh và metadata timestamp."),
        ("6", "Mô hình toán học calibration", "Diễn giải công thức dùng trong code."),
        ("6.1", "Chuyển đổi polar sang Cartesian", "x = d sin(theta), z = d cos(theta), y = 0."),
        ("6.2", "Tập điểm LiDAR 3D", "Biểu diễn điểm LiDAR dạng [x, 0, z]^T."),
        ("6.3", "Tập điểm ảnh 2D", "Pixel [u, v] lấy từ mapped_points trong JSON."),
        ("6.4", "Ước lượng ngoại tại", "camera = R * lidar + T."),
        ("6.5", "Chiếu điểm lên ảnh", "u = fx Xc/Zc + cx, v = fy Yc/Zc + cy."),
        ("7", "Giải thuật hiệu chuẩn bằng PnP/RANSAC", "Trình bày cụ thể thuật toán trong compute_calibration.py."),
        ("7.1", "Lấy camera intrinsics", "Đọc fx, fy, ppx, ppy, distortion từ RealSense SDK."),
        ("7.2", "Nạp dữ liệu pair JSON", "Duyệt các file pair_*.json và trích mapped_points."),
        ("7.3", "Thiết lập solvePnPRansac", "Flag, số vòng lặp, ngưỡng reprojectionError, confidence."),
        ("7.4", "Tính R và T", "Chuyển rvec sang R bằng Rodrigues, reshape T."),
        ("7.5", "Lưu kết quả calibration", "Lưu K, R, T vào calibration_result_pnp.npz."),
        ("8", "Triển khai chương trình", "Mô tả theo từng module code."),
        ("8.1", "Module thu dữ liệu", "Các hàm init_lidar, init_realsense, denoise_lidar_scans, calibrate."),
        ("8.2", "Module tính calibration", "Các hàm get_camera_intrinsics, load_data, solve_pnp, compute_error."),
        ("8.3", "Module fusion realtime", "Load K/R/T, đọc scan-frame, project và vẽ overlay."),
        ("8.4", "Cơ chế lưu kết quả", "Ảnh màu, ảnh depth, ảnh fusion, dữ liệu JSON/NPZ."),
        ("8.5", "Xử lý lỗi và dọn tài nguyên", "Stop LiDAR, disconnect, stop pipeline, destroy windows."),
        ("9", "Thiết kế thực nghiệm", "Cách tiến hành đo và tiêu chí đánh giá."),
        ("9.1", "Kịch bản thu dữ liệu", "Số mẫu, khoảng cách target, vị trí trái/phải/trung tâm."),
        ("9.2", "Điều kiện môi trường", "Ánh sáng, nền, vật cản, độ ổn định thiết bị."),
        ("9.3", "Tiêu chí chất lượng dữ liệu", "Số điểm mapped_points, outlier, độ rõ target."),
        ("9.4", "Chỉ số đánh giá", "Mean, median, max reprojection error, số inlier."),
        ("10", "Kết quả và đánh giá", "Phần để điền số liệu thực nghiệm sau khi chạy calibration."),
        ("10.1", "Kết quả camera intrinsics", "Bảng K và distortion coefficients."),
        ("10.2", "Kết quả ngoại tại R, T", "Ma trận quay, vector tịnh tiến, diễn giải vị trí tương đối."),
        ("10.3", "Thống kê inlier/outlier", "Số lượng correspondence và tỷ lệ inlier."),
        ("10.4", "Sai số tái chiếu", "Bảng mean/median/max và biểu đồ nếu cần."),
        ("10.5", "Minh họa kết quả projection", "Ảnh overlay điểm LiDAR lên ảnh RGB."),
        ("11", "Kiểm chứng bằng fusion realtime", "Dùng calibration_result_pnp.npz để kiểm tra trực quan."),
        ("11.1", "Quy trình chạy fusion", "Luồng xử lý trong fusion_calibration.py."),
        ("11.2", "Hiển thị RGB-depth", "Ghép overlay RGB và depth colormap."),
        ("11.3", "Đánh giá định tính", "Mức độ trùng khớp của điểm LiDAR với vật thể trong ảnh."),
        ("11.4", "Lưu ảnh kiểm chứng", "Ảnh color, depth, fusion trong data/fusion_output."),
        ("12", "Thảo luận", "Phân tích hạn chế và nguồn sai số."),
        ("12.1", "Nguồn sai số từ LiDAR", "Nhiễu khoảng cách, độ phân giải góc, phản xạ bề mặt."),
        ("12.2", "Nguồn sai số từ camera", "Sai số intrinsics, distortion, alignment depth-color."),
        ("12.3", "Nguồn sai số thao tác chọn điểm", "Sai lệch click hai đầu target và nội suy tuyến tính."),
        ("12.4", "Giới hạn của LiDAR 2D", "Giả định y = 0, target nằm trên mặt phẳng quét."),
        ("12.5", "Đề xuất cải tiến", "Tự động detect target, dùng nhiều target, lọc outlier tốt hơn."),
        ("13", "Kết luận", "Tổng kết pipeline, kết quả đạt được và hướng phát triển."),
        ("13.1", "Kết luận chính", "Tóm tắt quá trình calibration và fusion."),
        ("13.2", "Hướng phát triển", "Tự động hóa, đánh giá định lượng sâu hơn, mở rộng sang LiDAR 3D."),
        ("14", "Tài liệu tham khảo", "OpenCV PnP, RealSense SDK, RPLidar SDK, tài liệu calibration camera."),
        ("15", "Phụ lục", "Code, cấu hình chạy, mẫu JSON, bảng dữ liệu, lỗi thường gặp."),
        ("15.1", "Phụ lục A: Cấu trúc file JSON", "Giải thích mapped_points và cartesian_points."),
        ("15.2", "Phụ lục B: Hướng dẫn chạy chương trình", "Lệnh chạy collect, compute, fusion."),
        ("15.3", "Phụ lục C: Checklist phần cứng", "COM port, USB 3.0, nguồn, vị trí target."),
    ]

    toc_table = doc.add_table(rows=1, cols=3)
    toc_table.style = "Table Grid"
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Mục", "Tên mục", "Nội dung cần triển khai"]
    for idx, header in enumerate(headers):
        set_cell_text(toc_table.rows[0].cells[idx], header, bold=True)
        shade_cell(toc_table.rows[0].cells[idx], "D9EAD3")

    for number, title_text, description in toc:
        cells = toc_table.add_row().cells
        set_cell_text(cells[0], number, bold=number.count(".") == 0)
        set_cell_text(cells[1], title_text, bold=number.count(".") == 0)
        set_cell_text(cells[2], description)

    add_heading(doc, "GỢI Ý PHÂN BỔ TRANG", 1)
    allocation = [
        ("Giới thiệu và cơ sở lý thuyết", "3-5 trang"),
        ("Phần cứng, phần mềm và dữ liệu", "3-4 trang"),
        ("Phương pháp calibration PnP/RANSAC", "4-6 trang"),
        ("Thực nghiệm, kết quả và đánh giá", "5-8 trang"),
        ("Thảo luận, kết luận, phụ lục", "3-5 trang"),
    ]
    alloc_table = doc.add_table(rows=1, cols=2)
    alloc_table.style = "Table Grid"
    set_cell_text(alloc_table.rows[0].cells[0], "Nhóm nội dung", bold=True)
    set_cell_text(alloc_table.rows[0].cells[1], "Dung lượng đề xuất", bold=True)
    shade_cell(alloc_table.rows[0].cells[0], "FCE5CD")
    shade_cell(alloc_table.rows[0].cells[1], "FCE5CD")
    for left, right in allocation:
        row = alloc_table.add_row().cells
        set_cell_text(row[0], left)
        set_cell_text(row[1], right)

    doc.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
